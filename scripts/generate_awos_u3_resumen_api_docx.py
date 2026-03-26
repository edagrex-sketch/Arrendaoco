from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_PATH = r"C:\webapps\laravel\Arrendaoco\docs\awos_u3_resumen_api_arrendaoco.docx"


document = Document()

style = document.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Resumen del API del proyecto ArrendaOco")
title_run.bold = True
title_run.font.size = Pt(14)

sections = [
    (
        "Lenguaje de programacion y framework utilizado para el desarrollo",
        "El proyecto ArrendaOco esta desarrollado con PHP 8.2 y el framework Laravel 12. "
        "Esta combinacion permite construir una API robusta para aplicaciones web orientadas "
        "a servicios, ya que Laravel integra enrutamiento, controladores, validacion, "
        "autenticacion, ORM y herramientas para responder en formato JSON. Ademas, el proyecto "
        "incluye paquetes como Laravel Sanctum para autenticacion con tokens, Laravel Reverb "
        "para eventos en tiempo real, DomPDF para generar documentos PDF y maatwebsite/excel "
        "para exportaciones. Esto muestra que el framework no solo sirve para crear vistas web, "
        "sino tambien para exponer servicios reutilizables por clientes web y moviles."
    ),
    (
        "Elementos que componen la aplicacion web orientada a servicios",
        "La aplicacion se compone de varios elementos organizados por capas. En la capa de rutas, "
        "el archivo routes/api.php define endpoints publicos y protegidos. En la capa de control, "
        "existen controladores especializados como AuthController, InmuebleController, "
        "ContratoController, PagoController, ReporteController y ChatController, cada uno "
        "responsable de una parte del negocio. En la capa de datos se encuentran modelos como "
        "Usuario, Inmueble, Contrato, Pago, Resena, Evento, Chat y Mensaje, respaldados por "
        "migraciones que estructuran la base de datos. Tambien se utilizan recursos como "
        "InmuebleResource y UserResource para transformar la informacion en respuestas JSON "
        "limpias, asi como politicas y middleware para controlar acceso y seguridad."
    ),
    (
        "APIs propias desarrolladas y explicacion de su funcionamiento",
        "El proyecto desarrolla varias APIs propias. La API de autenticacion permite registrar "
        "usuarios, iniciar sesion, cerrar sesion y consultar el perfil autenticado mediante "
        "tokens. La API de inmuebles ofrece listado publico, detalle, alta, actualizacion y "
        "eliminacion de propiedades, ademas de filtros y carga de imagenes. La API de contratos "
        "gestiona la renta de un inmueble, la consulta de contratos y el estado de cuenta, "
        "incluyendo generacion de documentos. La API de pagos crea mensualidades, lista adeudos "
        "y registra pagos realizados por el inquilino. La API de reportes resume ingresos por "
        "mes y anio para apoyar la administracion. La API de favoritos y resenas permite guardar "
        "inmuebles de interes y dejar opiniones. Finalmente, la API de chat administra "
        "conversaciones entre usuarios y usa eventos para enviar mensajes en tiempo real. En "
        "conjunto, estas APIs convierten a ArrendaOco en una aplicacion orientada a servicios "
        "porque la logica del negocio queda disponible para distintos clientes sin duplicar procesos."
    ),
]

for heading, body in sections:
    p_heading = document.add_paragraph()
    run_heading = p_heading.add_run(heading)
    run_heading.bold = True
    run_heading.font.size = Pt(12)

    p_body = document.add_paragraph(body)
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.save(OUTPUT_PATH)
print(f"DOCX generado en: {OUTPUT_PATH}")
